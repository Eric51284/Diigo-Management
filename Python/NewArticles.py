import pandas as pd
from docx import Document
import requests
from bs4 import BeautifulSoup
import re
from datetime import datetime
import time
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from urllib.parse import urljoin, urlparse
import logging
import os
import json

# Set up logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


def extract_articles_and_links(docx_path):
    """Extract article titles and hyperlinks from Word document"""
    doc = Document(docx_path)
    articles = []

    logger.info(f"Reading document: {docx_path}")
    logger.info(f"Total paragraphs in document: {len(doc.paragraphs)}")

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()

        if not text:
            continue

        # More flexible matching - same as debug version
        if text.startswith("-") or any(
            title_word in text.lower()
            for title_word in ["microsoft", "ai", "anthropic", "openai", "chatgpt"]
        ):
            # Clean up title (remove leading dash and extra spaces)
            title = text.lstrip("- ").strip()

            # Try multiple methods to extract hyperlink
            hyperlink_url = (
                extract_hyperlink_method1(paragraph)
                or extract_hyperlink_method2(paragraph)
                or extract_hyperlink_method3(paragraph)
            )

            print(f"Found article {len(articles)+1}: {title[:60]}...")
            print(f"  URL: {hyperlink_url}")

            articles.append(
                {
                    "title": title,
                    "url": hyperlink_url,
                    "publication_date": None,
                    "status": "pending" if hyperlink_url else "no_url",
                }
            )

    logger.info(f"Found {len(articles)} articles")
    return articles


def extract_hyperlink_method1(paragraph):
    """Method 1: Standard hyperlink extraction"""
    try:
        for elem in paragraph._element.iter():
            if "hyperlink" in str(elem.tag).lower():
                r_id = elem.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                )
                if (
                    r_id
                    and hasattr(paragraph.part, "rels")
                    and r_id in paragraph.part.rels
                ):
                    url = paragraph.part.rels[r_id].target_ref
                    return url
    except Exception as e:
        pass
    return None


def extract_hyperlink_method2(paragraph):
    """Method 2: Check runs for hyperlinks"""
    try:
        for run in paragraph.runs:
            # Check if this run has hyperlink formatting
            if run._element.rPr is not None:
                for elem in run._element.iter():
                    if "hyperlink" in str(elem.tag).lower():
                        # Try to get the relationship ID
                        r_id = elem.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                        )
                        if (
                            r_id
                            and hasattr(paragraph.part, "rels")
                            and r_id in paragraph.part.rels
                        ):
                            url = paragraph.part.rels[r_id].target_ref
                            return url
    except Exception as e:
        pass
    return None


def extract_hyperlink_method3(paragraph):
    """Method 3: Deep dive into XML structure"""
    try:
        # Get the raw XML of the paragraph
        xml_str = str(paragraph._element.xml)

        # Look for hyperlink relationships
        hyperlink_pattern = r'r:id="(rId\d+)"'
        matches = re.findall(hyperlink_pattern, xml_str)

        for r_id in matches:
            if hasattr(paragraph.part, "rels") and r_id in paragraph.part.rels:
                url = paragraph.part.rels[r_id].target_ref
                return url

    except Exception as e:
        pass
    return None


def process_articles_for_dates(articles, delay=2):
    """Process all articles to get publication dates"""
    articles_with_urls = [a for a in articles if a["url"]]
    total = len(articles_with_urls)
    processed = 0

    logger.info(f"Processing {total} articles with URLs...")

    for article in articles:
        if not article["url"]:
            continue

        processed += 1
        logger.info(f"Processing {processed}/{total}: {article['title'][:50]}...")

        date, status = get_publication_date_from_url(article["url"], article["title"])
        article["publication_date"] = date
        article["status"] = status

        # Add delay to be respectful to servers
        if processed < total:
            time.sleep(delay)

    return articles


def get_publication_date_from_url(url, title):
    """Visit URL and extract publication date"""
    try:
        logger.info(f"Fetching: {url[:100]}...")

        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }

        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, "html.parser")

        # Try multiple methods to find publication date
        date = (
            find_date_in_meta_tags(soup)
            or find_date_in_json_ld(soup)
            or find_date_in_time_tags(soup)
            or find_date_in_article_tags(soup)
            or find_date_in_text_patterns(soup)
        )

        if date:
            logger.info(f"Found date: {date}")
            return date, "success"
        else:
            logger.warning(f"No date found for: {title[:50]}")
            return None, "no_date_found"

    except requests.exceptions.Timeout:
        logger.error(f"Timeout for: {url}")
        return None, "timeout"
    except requests.exceptions.RequestException as e:
        logger.error(f"Request error for {url}: {e}")
        return None, "request_error"
    except Exception as e:
        logger.error(f"Unexpected error for {url}: {e}")
        return None, "error"


def find_date_in_meta_tags(soup):
    """Find date in meta tags"""
    meta_selectors = [
        'meta[property="article:published_time"]',
        'meta[property="article:published"]',
        'meta[name="publish-date"]',
        'meta[name="publication-date"]',
        'meta[name="date"]',
        'meta[name="DC.date"]',
        'meta[name="DC.Date"]',
        'meta[property="og:published_time"]',
        'meta[name="publishdate"]',
        'meta[name="pub_date"]',
        'meta[itemprop="datePublished"]',
        'meta[itemprop="publishDate"]',
    ]

    for selector in meta_selectors:
        meta = soup.select_one(selector)
        if meta:
            content = meta.get("content") or meta.get("value")
            if content:
                parsed_date = parse_date_string(content)
                if parsed_date:
                    return parsed_date

    return None


def find_date_in_json_ld(soup):
    """Find date in JSON-LD structured data"""
    try:
        scripts = soup.find_all("script", type="application/ld+json")
        for script in scripts:
            if script.string:
                data = json.loads(script.string)

                # Handle both single objects and arrays
                if isinstance(data, list):
                    for item in data:
                        date = extract_date_from_json_object(item)
                        if date:
                            return date
                else:
                    date = extract_date_from_json_object(data)
                    if date:
                        return date
    except:
        pass

    return None


def extract_date_from_json_object(obj):
    """Extract date from JSON-LD object"""
    if not isinstance(obj, dict):
        return None

    date_fields = ["datePublished", "publishDate", "dateCreated", "uploadDate"]
    for field in date_fields:
        if field in obj:
            parsed_date = parse_date_string(obj[field])
            if parsed_date:
                return parsed_date

    return None


def find_date_in_time_tags(soup):
    """Find date in time tags"""
    time_selectors = [
        "time[datetime]",
        "time[pubdate]",
        ".published-date time",
        ".publish-date time",
        ".date time",
    ]

    for selector in time_selectors:
        time_elem = soup.select_one(selector)
        if time_elem:
            datetime_attr = time_elem.get("datetime") or time_elem.get("pubdate")
            if datetime_attr:
                parsed_date = parse_date_string(datetime_attr)
                if parsed_date:
                    return parsed_date

            # Try text content if no datetime attribute
            text = time_elem.get_text().strip()
            if text:
                parsed_date = parse_date_string(text)
                if parsed_date:
                    return parsed_date

    return None


def find_date_in_article_tags(soup):
    """Find date in article-related elements"""
    date_selectors = [
        ".published-date",
        ".publish-date",
        ".publication-date",
        ".date-published",
        ".article-date",
        ".post-date",
        ".entry-date",
        ".timestamp",
        '[class*="date"]',
        '[class*="publish"]',
    ]

    for selector in date_selectors:
        elements = soup.select(selector)
        for elem in elements:
            text = elem.get_text().strip()
            if text and len(text) < 100:  # Reasonable length for a date
                parsed_date = parse_date_string(text)
                if parsed_date:
                    return parsed_date

    return None


def find_date_in_text_patterns(soup):
    """Find date using text patterns in the page"""
    # Look for common date patterns in the HTML text
    text = soup.get_text()

    # Common date patterns
    patterns = [
        r"Published:?\s*([A-Za-z]+ \d{1,2},? \d{4})",
        r"Publication Date:?\s*([A-Za-z]+ \d{1,2},? \d{4})",
        r"(\d{1,2}/\d{1,2}/\d{4})",
        r"(\d{4}-\d{2}-\d{2})",
        r"([A-Za-z]+ \d{1,2},? \d{4})",
    ]

    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            parsed_date = parse_date_string(match)
            if parsed_date:
                return parsed_date

    return None


def parse_date_string(date_str):
    """Parse various date string formats into YYYY-MM-DD format"""
    if not date_str:
        return None

    # Clean the string
    date_str = str(date_str).strip()

    # Common date formats to try
    formats = [
        "%Y-%m-%d",
        "%Y-%m-%dT%H:%M:%S",
        "%Y-%m-%dT%H:%M:%SZ",
        "%Y-%m-%dT%H:%M:%S.%fZ",
        "%B %d, %Y",
        "%b %d, %Y",
        "%m/%d/%Y",
        "%d/%m/%Y",
        "%Y/%m/%d",
        "%B %d %Y",
        "%b %d %Y",
        "%d %B %Y",
        "%d %b %Y",
    ]

    for fmt in formats:
        try:
            dt = datetime.strptime(date_str, fmt)
            return dt.strftime("%Y-%m-%d")
        except ValueError:
            continue

    # Try to extract just the date part if it contains extra info
    date_match = re.search(r"(\d{4}-\d{2}-\d{2})", date_str)
    if date_match:
        return date_match.group(1)

    return None


def create_excel_file(articles, output_path):
    """Create Excel file with results"""
    # Make sure output directory exists
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    logger.info(f"Creating Excel file: {output_path}")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Article Publication Dates"

    # Headers
    headers = ["Title", "Publication Date", "URL", "Status"]

    # Add headers with formatting
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(
            start_color="366092", end_color="366092", fill_type="solid"
        )
        cell.alignment = Alignment(horizontal="center")

    # Add data
    for row_idx, article in enumerate(articles, 2):
        # Title with hyperlink
        title_cell = ws.cell(row=row_idx, column=1, value=article["title"])
        if article["url"]:
            title_cell.hyperlink = article["url"]
            title_cell.font = Font(color="0563C1", underline="single")

        # Publication date
        ws.cell(row=row_idx, column=2, value=article["publication_date"])

        # URL
        ws.cell(row=row_idx, column=3, value=article["url"])

        # Status
        status_cell = ws.cell(row=row_idx, column=4, value=article["status"])

        # Color code status
        if article["status"] == "success":
            status_cell.fill = PatternFill(
                start_color="C6EFCE", end_color="C6EFCE", fill_type="solid"
            )
        elif "error" in article["status"] or "timeout" in article["status"]:
            status_cell.fill = PatternFill(
                start_color="FFC7CE", end_color="FFC7CE", fill_type="solid"
            )
        elif article["status"] == "no_date_found":
            status_cell.fill = PatternFill(
                start_color="FFEB9C", end_color="FFEB9C", fill_type="solid"
            )

    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 80)
        ws.column_dimensions[column_letter].width = adjusted_width

    wb.save(output_path)
    logger.info(f"Excel file saved successfully")


def main(docx_path=None, output_path=None):
    """Main function"""
    # Corrected relative paths
    # docx_path = r".\Source files\Diigo New Articles.docx"
    # output_path = r".\Output files\Article_Publication_Dates.xlsx"

    try:
        # Step 1: Extract articles and links from Word document
        articles = extract_articles_and_links(docx_path)

        if not articles:
            logger.error("No articles found in the document")
            return

        # Step 2: Process articles to get publication dates
        processed_articles = process_articles_for_dates(articles, delay=2)

        # Step 3: Create Excel file
        create_excel_file(processed_articles, output_path)

        # Summary
        total = len(processed_articles)
        successful = sum(1 for a in processed_articles if a["status"] == "success")
        no_date = sum(1 for a in processed_articles if a["status"] == "no_date_found")
        errors = sum(
            1
            for a in processed_articles
            if "error" in a["status"] or "timeout" in a["status"]
        )
        no_url = sum(1 for a in processed_articles if a["status"] == "no_url")

        print(f"\n=== SUMMARY ===")
        print(f"Total articles: {total}")
        print(f"Successfully found dates: {successful}")
        print(f"No date found: {no_date}")
        print(f"Errors/timeouts: {errors}")
        print(f"No URL: {no_url}")
        print(f"\nResults saved to: {os.path.abspath(output_path)}")

    except Exception as e:
        logger.error(f"Error in main process: {e}")
        raise


if __name__ == "__main__":
    # Set default paths here if running as a script
    default_docx_path = r".\Source files\Diigo New Articles.docx"
    default_output_path = r".\Output files\Article_Publication_Dates.xlsx"
    main(default_docx_path, default_output_path)
