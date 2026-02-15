from docx import Document
import re


def analyze_word_document(docx_path):
    """Analyze the Word document to see all content and identify potential articles"""
    print(f"Analyzing document: {docx_path}")

    doc = Document(docx_path)

    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print("=" * 80)

    # Current detection logic
    articles_found_current = []

    # All potential articles (broader detection)
    all_potential_articles = []

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()

        if not text:
            continue

        # Show every non-empty paragraph with its index
        print(f"Paragraph {i+1:3d}: {repr(text[:100])}...")

        # Current detection logic (what the program currently finds)
        if text.startswith("-") or any(
            keyword in text.lower()
            for keyword in ["microsoft", "ai", "anthropic", "openai", "chatgpt"]
        ):
            articles_found_current.append(
                {
                    "index": i + 1,
                    "text": text,
                    "reason": (
                        "starts_with_dash"
                        if text.startswith("-")
                        else "contains_ai_keyword"
                    ),
                }
            )

        # Broader detection - anything that looks like it could be an article
        if (
            len(text) > 10  # Reasonable length
            and not text.isupper()  # Not a header
            and not text.startswith("Page ")  # Not page number
            and (
                "http" in text  # Contains URL
                or any(
                    word in text.lower()
                    for word in [
                        "article",
                        "report",
                        "study",
                        "news",
                        "company",
                        "technology",
                    ]
                )
                or text.startswith("-")
                or text.startswith("•")
                or text.startswith("*")
                or re.match(r"^\d+\.", text)  # Numbered list
                or any(
                    keyword in text.lower()
                    for keyword in [
                        "microsoft",
                        "ai",
                        "anthropic",
                        "openai",
                        "chatgpt",
                        "google",
                        "amazon",
                        "meta",
                        "tesla",
                    ]
                )
            )
        ):
            all_potential_articles.append({"index": i + 1, "text": text})

    print("\n" + "=" * 80)
    print(f"CURRENT DETECTION RESULTS:")
    print(f"Found {len(articles_found_current)} articles with current logic:")

    for article in articles_found_current:
        print(
            f"  {article['index']:3d}: {article['text'][:60]}... ({article['reason']})"
        )

    print("\n" + "=" * 80)
    print(f"BROADER DETECTION RESULTS:")
    print(f"Found {len(all_potential_articles)} potential articles:")

    for article in all_potential_articles:
        is_current = article["index"] in [a["index"] for a in articles_found_current]
        marker = "✓" if is_current else "✗"
        print(f"  {marker} {article['index']:3d}: {article['text'][:60]}...")

    print("\n" + "=" * 80)
    print("MISSED ARTICLES (broader detection found but current logic missed):")

    current_indices = [a["index"] for a in articles_found_current]
    missed_articles = [
        a for a in all_potential_articles if a["index"] not in current_indices
    ]

    for article in missed_articles:
        print(f"  MISSED {article['index']:3d}: {article['text']}")
        print(
            f"    Reason it was missed: Does not start with '-' and doesn't contain AI keywords"
        )
        print()

    return articles_found_current, all_potential_articles, missed_articles


def suggest_improved_detection(missed_articles):
    """Suggest improvements to the detection logic based on missed articles"""
    if not missed_articles:
        print("No missed articles - current detection is perfect!")
        return

    print("SUGGESTED IMPROVEMENTS TO DETECTION LOGIC:")
    print("=" * 50)

    # Analyze patterns in missed articles
    patterns = {
        "starts_with_bullet": [],
        "starts_with_number": [],
        "contains_company_names": [],
        "contains_tech_terms": [],
        "other": [],
    }

    for article in missed_articles:
        text = article["text"]

        if text.startswith("•") or text.startswith("*"):
            patterns["starts_with_bullet"].append(text)
        elif re.match(r"^\d+\.", text):
            patterns["starts_with_number"].append(text)
        elif any(
            company in text.lower()
            for company in [
                "apple",
                "google",
                "amazon",
                "meta",
                "tesla",
                "nvidia",
                "intel",
            ]
        ):
            patterns["contains_company_names"].append(text)
        elif any(
            term in text.lower()
            for term in ["technology", "software", "tech", "digital", "innovation"]
        ):
            patterns["contains_tech_terms"].append(text)
        else:
            patterns["other"].append(text)

    for pattern_name, articles in patterns.items():
        if articles:
            print(f"\n{pattern_name.upper()} ({len(articles)} articles):")
            for article in articles[:3]:  # Show first 3 examples
                print(f"  Example: {article[:80]}...")


def main():
    """Main function"""
    docx_path = r".\Output files\Diigo Outlined Articles.docx"

    try:
        current_articles, all_potential, missed = analyze_word_document(docx_path)
        suggest_improved_detection(missed)

        print(f"\n" + "=" * 80)
        print("SUMMARY:")
        print(f"Current logic finds: {len(current_articles)} articles")
        print(f"Broader detection finds: {len(all_potential)} potential articles")
        print(f"You said there are: 194 articles")
        print(f"Missing: {194 - len(current_articles)} articles")

        if missed:
            print(f"\nTo find all 194 articles, you may need to:")
            print("1. Update the detection logic to include the missed patterns")
            print("2. Or manually add the missing articles to your Excel file")

    except Exception as e:
        print(f"Error: {e}")
        raise


if __name__ == "__main__":
    main()
