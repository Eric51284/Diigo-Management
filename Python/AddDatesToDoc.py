import pandas as pd
from docx import Document
import openpyxl
import os
import re
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy


def read_dates_from_excel(excel_path):
    """Read the publication dates from the Excel file"""
    print(f"Reading dates from: {excel_path}")

    # Read the Excel file
    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active

    # Create a dictionary to map titles to dates
    title_to_date = {}

    # Skip header row and read data
    for row in ws.iter_rows(min_row=2, values_only=True):
        title = row[0]  # Column A: Title
        pub_date = row[1]  # Column B: Publication Date

        if title and pub_date:
            # Clean the title for matching
            clean_title = title.strip()
            title_to_date[clean_title] = str(pub_date)

    print(f"Found dates for {len(title_to_date)} articles")
    return title_to_date


def update_word_document_with_dates(docx_path, title_to_date, output_path):
    """Update the Word document by prepending dates to article titles while preserving hyperlinks"""
    print(f"Reading original document: {docx_path}")

    # Open the original document
    doc = Document(docx_path)

    # Track statistics
    updated_count = 0
    not_found_count = 0

    print(f"Processing {len(doc.paragraphs)} paragraphs...")

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()

        if not text:
            continue

        # Check if this looks like an article entry
        if text.startswith("-") or any(
            keyword in text.lower()
            for keyword in ["microsoft", "ai", "anthropic", "openai", "chatgpt"]
        ):
            # Clean up the title to match against our Excel data
            original_title = text.lstrip("- ").strip()

            # Look for a matching date
            matching_date = None
            for excel_title, date in title_to_date.items():
                if titles_match(original_title, excel_title):
                    matching_date = date
                    break

            if matching_date:
                try:
                    success = prepend_date_preserve_hyperlinks(
                        paragraph, matching_date, text.startswith("-")
                    )
                    if success:
                        updated_count += 1
                        print(f"Updated: {original_title[:50]}... -> {matching_date}")
                    else:
                        print(f"Failed to update: {original_title[:50]}...")
                except Exception as e:
                    print(f"Error updating paragraph {i+1}: {e}")
                    print(f"  Text: {text[:50]}...")
            else:
                not_found_count += 1
                print(f"No date found for: {original_title[:50]}...")

    # Save the updated document
    print(f"Saving updated document to: {output_path}")
    doc.save(output_path)

    print(f"\n=== UPDATE SUMMARY ===")
    print(f"Articles updated with dates: {updated_count}")
    print(f"Articles without dates: {not_found_count}")
    print(f"Updated document saved to: {os.path.abspath(output_path)}")

    return updated_count, not_found_count


def prepend_date_preserve_hyperlinks(paragraph, date, has_dash):
    """Prepend date while preserving hyperlinks by working with XML directly"""
    try:
        # Create the date prefix
        if has_dash:
            date_prefix = f"- {date} - "
        else:
            date_prefix = f"{date} - "

        # Get the paragraph XML element
        p_elem = paragraph._element

        # Create a new run element for the date
        new_run = OxmlElement("w:r")

        # Create text element
        new_text = OxmlElement("w:t")
        new_text.text = date_prefix
        new_run.append(new_text)

        # Insert at the beginning of the paragraph
        p_elem.insert(0, new_run)

        # If the original started with "- ", we need to remove it from the existing content
        if has_dash:
            remove_leading_dash_from_paragraph(p_elem)

        return True

    except Exception as e:
        print(f"Error in prepend_date_preserve_hyperlinks: {e}")
        return False


def remove_leading_dash_from_paragraph(p_elem):
    """Remove the leading '- ' from the paragraph content"""
    try:
        # Find the first text node and remove "- " from it
        for elem in p_elem.iter():
            if elem.tag.endswith("}t") and elem.text:  # This is a text element
                if elem.text.startswith("- "):
                    elem.text = elem.text[2:]  # Remove "- "
                    break
                elif elem.text.startswith("-"):
                    elem.text = elem.text[1:]  # Remove just "-"
                    break
    except Exception as e:
        print(f"Error removing leading dash: {e}")


def titles_match(title1, title2):
    """Check if two titles are similar enough to be considered a match"""
    # Clean both titles
    clean1 = clean_title_for_matching(title1)
    clean2 = clean_title_for_matching(title2)

    # Exact match
    if clean1 == clean2:
        return True

    # Check if one is contained in the other (for cases where title might be truncated)
    if len(clean1) > 20 and len(clean2) > 20:
        if clean1 in clean2 or clean2 in clean1:
            return True

    # Check first 50 characters
    if clean1[:50] == clean2[:50]:
        return True

    return False


def clean_title_for_matching(title):
    """Clean title for better matching"""
    # Convert to lowercase
    title = title.lower()

    # Remove extra whitespace
    title = re.sub(r"\s+", " ", title).strip()

    # Remove common punctuation that might differ
    title = re.sub(r'[""' '""' "]", '"', title)  # Normalize quotes
    title = re.sub(r"[–—]", "-", title)  # Normalize dashes

    return title


def main():
    """Main function to update Word document with dates"""
    # File paths
    excel_path = r".\Output files\Article_Publication_Dates.xlsx"
    original_docx = r".\Source files\Diigo New Articles.docx"
    output_docx = r".\Output files\Diigo New Articles with Dates.docx"

    try:
        # Check if files exist
        if not os.path.exists(excel_path):
            print(f"ERROR: Excel file not found: {excel_path}")
            return

        if not os.path.exists(original_docx):
            print(f"ERROR: Original Word document not found: {original_docx}")
            return

        # Make sure output directory exists
        output_dir = os.path.dirname(output_docx)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # Step 1: Read dates from Excel file
        title_to_date = read_dates_from_excel(excel_path)

        if not title_to_date:
            print("No dates found in Excel file!")
            return

        # Step 2: Update Word document with dates
        updated_count, not_found_count = update_word_document_with_dates(
            original_docx, title_to_date, output_docx
        )

        print(f"\nProcess completed successfully!")
        print(f"Original document: {os.path.abspath(original_docx)}")
        print(f"Updated document: {os.path.abspath(output_docx)}")
        print(f"Hyperlinks have been preserved!")

    except Exception as e:
        print(f"Error: {e}")
        raise


if __name__ == "__main__":
    main()
